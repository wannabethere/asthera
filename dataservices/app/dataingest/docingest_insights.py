import uuid
import json
import logging
import pickle
import re
from dataclasses import dataclass, asdict
from datetime import datetime
from typing import Any, Dict, List, Optional, Union, Tuple
from pathlib import Path
import hashlib

# External libraries
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document as LangchainDocument
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
import chromadb

# Internal imports (these would be your actual imports)
from app.storage.documents import DocumentChromaStore
from app.core.dependencies import build_document_store
from app.schemas.docs.docmodels import Document, DocumentInsight
from app.core.session_manager import SessionManager
from app.utils.history import DomainManager
from app.core.settings import get_settings,ServiceConfig
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

logger = logging.getLogger(__name__)
settings = get_settings()

@dataclass
class DocumentChunk:
    """Represents a chunk of a document for processing"""
    chunk_id: str
    content: str
    metadata: Dict[str, Any]


@dataclass
class ProcessingConfig:
    """Configuration for document processing"""
    chunk_size: int = 2000
    chunk_overlap: int = 200
    use_questions: bool = True
    enable_caching: bool = True
    store_in_db: bool = True
    store_in_chromadb: bool = True
    model_name: str = "gpt-4o"
    temperature: float = 0.0
    # Flexible extraction configuration
    extraction_types: List[str] = None  # List of extraction types to perform
    custom_extraction_config: Dict = None  # Custom configuration for extractions


@dataclass
class ExtractionConfig:
    """Configuration for different extraction types"""
    extraction_type: str
    enabled: bool = True
    priority: int = 1  # Higher number = higher priority
    custom_prompt: Optional[str] = None
    output_format: str = "json"  # json, text, structured
    metadata: Dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


# Default extraction configurations
DEFAULT_EXTRACTION_CONFIGS = {
    "business_intelligence": ExtractionConfig(
        extraction_type="business_intelligence",
        enabled=True,
        priority=1,
        custom_prompt=None,
        output_format="json",
        metadata={"categories": ["kpis", "business_terms", "data_definitions", "operational_metrics"]}
    ),
    "entities": ExtractionConfig(
        extraction_type="entities",
        enabled=True,
        priority=2,
        custom_prompt=None,
        output_format="json",
        metadata={"types": ["people", "organizations", "locations", "concepts"]}
    ),
    "financial_metrics": ExtractionConfig(
        extraction_type="financial_metrics",
        enabled=False,
        priority=3,
        custom_prompt="Extract financial metrics, ratios, and performance indicators",
        output_format="json",
        metadata={"focus": "financial_data"}
    ),
    "compliance_terms": ExtractionConfig(
        extraction_type="compliance_terms",
        enabled=False,
        priority=4,
        custom_prompt="Extract compliance-related terms, regulations, and requirements",
        output_format="json",
        metadata={"focus": "compliance"}
    )
}


class DocumentCache:
    """Simple file-based caching for processed documents"""
    
    def __init__(self, cache_dir: str = "./cache"):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        
    def _get_cache_key(self, content: str, config: ProcessingConfig) -> str:
        """Generate cache key from content and config"""
        content_hash = hashlib.md5(content.encode()).hexdigest()
        config_hash = hashlib.md5(str(asdict(config)).encode()).hexdigest()
        return f"{content_hash}_{config_hash}"
    
    def get(self, content: str, config: ProcessingConfig) -> Optional[Dict]:
        """Get cached result if exists"""
        try:
            cache_key = self._get_cache_key(content, config)
            cache_file = self.cache_dir / f"{cache_key}.pkl"
            
            if cache_file.exists():
                with open(cache_file, 'rb') as f:
                    return pickle.load(f)
        except Exception as e:
            logger.warning(f"Cache read error: {e}")
        return None
    
    def set(self, content: str, config: ProcessingConfig, result: Dict):
        """Cache the result"""
        try:
            cache_key = self._get_cache_key(content, config)
            cache_file = self.cache_dir / f"{cache_key}.pkl"
            
            with open(cache_file, 'wb') as f:
                pickle.dump(result, f)
        except Exception as e:
            logger.warning(f"Cache write error: {e}")


class InputProcessor:
    """Handles different input types and converts them to standardized format"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap
        )
    
    def process_pdf(self, pdf_path: str, metadata: Dict = None) -> Tuple[str, List[DocumentChunk]]:
        """Process PDF file and return content and chunks"""
        try:
            loader = PyPDFLoader(pdf_path)
            documents = loader.load()
            
            # Combine all pages
            full_content = "\n".join([doc.page_content for doc in documents])
            
            # Split into chunks
            chunks = self.text_splitter.split_documents(documents)
            
            # Convert to DocumentChunk objects
            doc_chunks = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    **(metadata or {}),
                    "chunk_index": i,
                    "source": pdf_path,
                    "page": chunk.metadata.get("page", 0)
                }
                
                doc_chunks.append(DocumentChunk(
                    chunk_id=f"chunk_{i}",
                    content=chunk.page_content,
                    metadata=chunk_metadata
                ))
            
            logger.info(f"Processed PDF: {len(documents)} pages, {len(doc_chunks)} chunks")
            return full_content, doc_chunks
            
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")
            raise
    
    def process_text(self, content: str, metadata: Dict = None) -> Tuple[str, List[DocumentChunk]]:
        """Process text content and return chunks"""
        try:
            # Create a temporary document for splitting
            doc = LangchainDocument(page_content=content, metadata=metadata or {})
            chunks = self.text_splitter.split_documents([doc])
            
            # Convert to DocumentChunk objects
            doc_chunks = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    **(metadata or {}),
                    "chunk_index": i,
                    "source": "text_input"
                }
                
                doc_chunks.append(DocumentChunk(
                    chunk_id=f"chunk_{i}",
                    content=chunk.page_content,
                    metadata=chunk_metadata
                ))
            
            logger.info(f"Processed text: {len(content)} chars, {len(doc_chunks)} chunks")
            return content, doc_chunks
            
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            raise
    
    def process_json(self, json_data: Dict, metadata: Dict = None) -> Tuple[str, List[DocumentChunk]]:
        """Process JSON data and return text representation and chunks"""
        try:
            # Convert JSON to readable text
            if "content" in json_data:
                content = str(json_data["content"])
            else:
                content = json.dumps(json_data, indent=2)
            
            # Merge metadata
            json_metadata = json_data.get("metadata", {})
            combined_metadata = {**(metadata or {}), **json_metadata}
            
            return self.process_text(content, combined_metadata)
            
        except Exception as e:
            logger.error(f"Error processing JSON: {e}")
            raise
    
    def process_docx(self, docx_path: str, metadata: Dict = None) -> Tuple[str, List[DocumentChunk]]:
        """Process Word document (.docx) and return content and chunks"""
        try:
            from docx import Document as DocxDocument
            
            # Load Word document
            doc = DocxDocument(docx_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            # Combine all content
            full_content = "\n".join(paragraphs)
            
            # Process as text
            combined_metadata = {
                **(metadata or {}),
                "source": docx_path,
                "document_type": "word_document",
                "file_extension": ".docx"
            }
            
            return self.process_text(full_content, combined_metadata)
            
        except ImportError:
            logger.error("python-docx library not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error processing Word document {docx_path}: {e}")
            raise
    
    def process_google_docs(self, doc_id: str, metadata: Dict = None) -> Tuple[str, List[DocumentChunk]]:
        """Process Google Docs document and return content and chunks"""
        try:
            from google.oauth2.credentials import Credentials
            from googleapiclient.discovery import build
            from googleapiclient.errors import HttpError
            
            # Initialize Google Docs API
            # Note: You'll need to set up OAuth2 credentials
            credentials = Credentials.from_authorized_user_file('token.json')
            service = build('docs', 'v1', credentials=credentials)
            
            # Get document content
            document = service.documents().get(documentId=doc_id).execute()
            
            # Extract text content
            content_parts = []
            for element in document.get('body', {}).get('content', []):
                if 'paragraph' in element:
                    paragraph = element['paragraph']
                    for text_run in paragraph.get('elements', []):
                        if 'textRun' in text_run:
                            content_parts.append(text_run['textRun']['content'])
                elif 'table' in element:
                    table = element['table']
                    for row in table.get('tableRows', []):
                        row_text = []
                        for cell in row.get('tableCells', []):
                            cell_content = []
                            for cell_element in cell.get('content', []):
                                if 'paragraph' in cell_element:
                                    for text_run in cell_element['paragraph'].get('elements', []):
                                        if 'textRun' in text_run:
                                            cell_content.append(text_run['textRun']['content'])
                            if cell_content:
                                row_text.append(''.join(cell_content).strip())
                        if row_text:
                            content_parts.append(' | '.join(row_text))
            
            full_content = ''.join(content_parts)
            
            # Process as text
            combined_metadata = {
                **(metadata or {}),
                "source": f"google_docs:{doc_id}",
                "document_type": "google_document",
                "google_doc_id": doc_id,
                "google_doc_title": document.get('title', '')
            }
            
            return self.process_text(full_content, combined_metadata)
            
        except ImportError:
            logger.error("Google API libraries not installed. Install with: pip install google-api-python-client google-auth-httplib2 google-auth-oauthlib")
            raise
        except HttpError as e:
            logger.error(f"Google Docs API error: {e}")
            raise
        except Exception as e:
            logger.error(f"Error processing Google Docs {doc_id}: {e}")
            raise
    
    def process_wiki(self, wiki_content: str, metadata: Dict = None) -> Tuple[str, List[DocumentChunk]]:
        """Process Wiki content (Markdown format) and return content and chunks"""
        try:
            import re
            
            # Clean up wiki markup
            content = self._clean_wiki_markup(wiki_content)
            
            # Process as text
            combined_metadata = {
                **(metadata or {}),
                "source": metadata.get("source", "wiki_input"),
                "document_type": "wiki_document",
                "content_format": "markdown"
            }
            
            return self.process_text(content, combined_metadata)
            
        except Exception as e:
            logger.error(f"Error processing Wiki content: {e}")
            raise
    
    def process_wiki_file(self, wiki_file_path: str, metadata: Dict = None) -> Tuple[str, List[DocumentChunk]]:
        """Process Wiki file and return content and chunks"""
        try:
            with open(wiki_file_path, 'r', encoding='utf-8') as f:
                wiki_content = f.read()
            
            file_metadata = {
                **(metadata or {}),
                "source": wiki_file_path,
                "file_extension": ".md" if wiki_file_path.endswith('.md') else ".txt"
            }
            
            return self.process_wiki(wiki_content, file_metadata)
            
        except Exception as e:
            logger.error(f"Error processing Wiki file {wiki_file_path}: {e}")
            raise
    
    def _clean_wiki_markup(self, content: str) -> str:
        """Clean up wiki markup and convert to readable text"""
        # Remove wiki markup patterns
        patterns = [
            (r'\[\[([^|\]]+)\|([^\]]+)\]\]', r'\2'),  # [[link|text]] -> text
            (r'\[\[([^\]]+)\]\]', r'\1'),  # [[link]] -> link
            (r'\{\{([^}]+)\}\}', ''),  # {{template}} -> empty
            (r'<ref[^>]*>.*?</ref>', ''),  # <ref>...</ref> -> empty
            (r'<ref[^>]*/>', ''),  # <ref/> -> empty
            (r'={2,}([^=]+)={2,}', r'\n\1\n'),  # ==Header== -> Header
            (r"'''([^']+)'''", r'\1'),  # '''bold''' -> bold
            (r"''([^']+)''", r'\1'),  # ''italic'' -> italic
            (r'^\*+', ''),  # Remove bullet points
            (r'^#+', ''),  # Remove numbered lists
            (r'^\|', ''),  # Remove table markers
            (r'^\s*$', ''),  # Remove empty lines
        ]
        
        cleaned_content = content
        for pattern, replacement in patterns:
            cleaned_content = re.sub(pattern, replacement, cleaned_content, flags=re.MULTILINE)
        
        # Clean up extra whitespace
        cleaned_content = re.sub(r'\n\s*\n', '\n\n', cleaned_content)
        cleaned_content = cleaned_content.strip()
        
        return cleaned_content
    


class AdvancedExtractor:
    """Advanced extraction using LLM-based analysis with flexible configuration"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        from app.core.dependencies import get_llm
        self.llm = get_llm()
        self.extraction_configs = self._load_extraction_configs()
        self.setup_extraction_chains()
    
    def _load_extraction_configs(self) -> Dict[str, ExtractionConfig]:
        """Load extraction configurations from config or use defaults"""
        if self.config.custom_extraction_config:
            # Merge custom config with defaults
            configs = DEFAULT_EXTRACTION_CONFIGS.copy()
            configs.update(self.config.custom_extraction_config)
            return configs
        return DEFAULT_EXTRACTION_CONFIGS
    
    def setup_extraction_chains(self):
        """Set up LangChain extraction chains based on enabled extraction types"""
        from langchain_core.output_parsers import JsonOutputParser
        from langchain_core.prompts import ChatPromptTemplate
        
        self.extraction_chains = {}
        
        # Get enabled extraction types
        enabled_types = self.config.extraction_types or ["business_intelligence", "entities"]
        
        for extraction_type in enabled_types:
            if extraction_type in self.extraction_configs:
                config = self.extraction_configs[extraction_type]
                if config.enabled:
                    self.extraction_chains[extraction_type] = self._create_extraction_chain(config)
    
    def _create_extraction_chain(self, config: ExtractionConfig):
        """Create extraction chain for a specific extraction type"""
        from langchain_core.output_parsers import JsonOutputParser
        from langchain_core.prompts import ChatPromptTemplate
        
        # Use custom prompt if provided, otherwise use default
        if config.custom_prompt:
            template = f"""
            {config.custom_prompt}
            
            Content: {{content}}
            
            Return as JSON with the requested information.
            """
        else:
            template = self._get_default_template(config.extraction_type)
        
        prompt = ChatPromptTemplate.from_template(template)
        parser = JsonOutputParser()
        return prompt | self.llm | parser
    
    def _get_default_template(self, extraction_type: str) -> str:
        """Get default template for extraction type"""
        templates = {
            "business_intelligence": """
            Extract comprehensive business intelligence from this document content.
            
            Content: {content}
            
            Extract the following categories:
            1. Key Performance Indicators (KPIs): Financial, operational, customer, strategic metrics
            2. Business Terms and Definitions: Industry terminology, acronyms, technical concepts
            3. Data Definitions and Schema: Data fields, types, sources, quality requirements
            4. Operational Metrics: Process efficiency, resource utilization, benchmarks
            5. Important Entities: People, organizations, products, departments
            6. Key Insights and Themes: Strategic insights, trends, recommendations
            
            Return as JSON with keys: kpis, business_terms, data_definitions, operational_metrics, entities, insights, ner_text
            """,
            "entities": """
            Extract named entities from this document content.
        
        Content: {content}
        
        Extract:
            1. People: Names, roles, titles
            2. Organizations: Companies, departments, teams
            3. Locations: Cities, countries, addresses
            4. Concepts: Key ideas, topics, themes
            5. Products: Services, technologies, solutions
            
            Return as JSON with keys: people, organizations, locations, concepts, products
            """,
            "financial_metrics": """
            Extract financial metrics and performance indicators from this document content.
            
            Content: {content}
            
            Extract:
            1. Revenue metrics: Sales, income, revenue streams
            2. Profit metrics: Net profit, gross profit, margins
            3. Growth metrics: Growth rates, trends, projections
            4. Financial ratios: ROI, ROE, debt ratios, liquidity ratios
            5. Budget information: Budgets, forecasts, allocations
            
            Return as JSON with keys: revenue_metrics, profit_metrics, growth_metrics, financial_ratios, budget_info
            """,
            "compliance_terms": """
            Extract compliance-related terms, regulations, and requirements from this document content.
        
        Content: {content}
            
            Extract:
            1. Regulations: Laws, rules, standards, frameworks
            2. Requirements: Mandatory requirements, obligations
            3. Compliance terms: Industry-specific compliance terminology
            4. Risk factors: Compliance risks, violations, penalties
            5. Controls: Internal controls, procedures, safeguards
            
            Return as JSON with keys: regulations, requirements, compliance_terms, risk_factors, controls
            """
        }
        return templates.get(extraction_type, templates["business_intelligence"])
    
    def extract_insights(self, content: str, questions: List[str] = None) -> Dict:
        """Extract insights using configured extraction types"""
        try:
            all_insights = {}
            extraction_metadata = {}
            
            # Run each enabled extraction type
            for extraction_type, chain in self.extraction_chains.items():
                try:
                    logger.info(f"Running {extraction_type} extraction")
                    
                    if questions and self.config.use_questions:
                        # Use question-based extraction if available
                        questions_text = "\n".join([f"- {q}" for q in questions])
                        result = chain.invoke({
                            "content": content[:8000],
                            "questions": questions_text
                        })
                    else:
                        # Use basic extraction
                        result = chain.invoke({"content": content[:8000]})
                    
                    # Store results under extraction type
                    all_insights[extraction_type] = result
                    
                    # Store metadata about this extraction
                    config = self.extraction_configs.get(extraction_type)
                    extraction_metadata[extraction_type] = {
                        "enabled": config.enabled if config else False,
                        "priority": config.priority if config else 1,
                        "output_format": config.output_format if config else "json",
                        "metadata": config.metadata if config else {}
                    }
                    
                except Exception as e:
                    logger.error(f"Error in {extraction_type} extraction: {e}")
                    all_insights[extraction_type] = {}
            
            # Add legacy fields for backward compatibility
            if "business_intelligence" in all_insights:
                bi_result = all_insights["business_intelligence"]
                all_insights["phrases"] = bi_result.get("phrases", [])
                all_insights["insights"] = bi_result.get("insights", {})
                all_insights["ner_text"] = bi_result.get("ner_text", "")
            
            return {
                "insights": all_insights,
                "extraction_metadata": extraction_metadata,
                "extraction_types": list(self.extraction_chains.keys())
            }
            
        except Exception as e:
            logger.error(f"Error in insight extraction: {e}")
            return {
                "insights": {},
                "extraction_metadata": {},
                "extraction_types": []
            }
    
    def extract_basic_insights(self, content: str) -> Dict:
        """Extract insights without questions (backward compatibility)"""
        return self.extract_insights(content, questions=None)
    
    def extract_with_questions(self, content: str, questions: List[str]) -> Dict:
        """Extract insights based on specific questions (backward compatibility)"""
        return self.extract_insights(content, questions=questions)


class DocumentIngestionService:
    """Main service for document ingestion and processing using session manager pattern"""
    
    def __init__(
        self,
        session_manager: SessionManager,
        domain_manager: DomainManager,
        config: ProcessingConfig = None,
        chroma_client: Optional[chromadb.PersistentClient] = None,
        document_store: Any = None,
    ):
        self.session_manager = session_manager
        self.domain_manager = domain_manager
        self.config = config or ProcessingConfig()
        self.cache = DocumentCache() if self.config.enable_caching else None
        self.input_processor = InputProcessor(self.config)
        self.extractor = AdvancedExtractor(self.config)

        if document_store is not None:
            self.chroma_store = document_store
        elif chroma_client is not None:
            self.chroma_store = DocumentChromaStore(
                client=chroma_client,
                collection_name="documents",
                tf_idf=True,
            )
        elif self.config.store_in_chromadb:
            self.chroma_store = build_document_store("documents", tf_idf=True)
        else:
            self.chroma_store = None

        logger.info("Document Ingestion Service initialized with session manager")
        
    async def ingest_document(self,
                       input_data: Union[str, Dict],
                             input_type: str,  # "pdf", "text", "json", "docx", "google_docs", "wiki", "wiki_file"
                       source_type: str,
                       document_type: str,
                       created_by: str,
                             domain_id: str,
                       questions: Optional[List[str]] = None,
                       metadata: Optional[Dict] = None,
                       event_type: str = "document_ingestion") -> Tuple[Document, DocumentInsight]:
        """
        Main ingestion method that handles any input type using session manager
        
        Args:
            input_data: File path (for PDF) or content (for text/json)
            input_type: Type of input ("pdf", "text", "json", "docx", "google_docs", "wiki", "wiki_file")
            source_type: Source of the document
            document_type: Type/category of document
            created_by: User/system that created the document
            domain_id: Domain ID for the document
            questions: Optional questions for targeted extraction
            metadata: Additional metadata
            event_type: Type of processing event
            
        Returns:
            Tuple of (Document, DocumentInsight)
        """
        async with self.session_manager.get_async_db_session() as session:
            try:
                start_time = datetime.now()
                logger.info(f"Starting ingestion: {input_type} document from {source_type}")
                
                # Process input based on type
                if input_type == "pdf":
                    content, chunks = self.input_processor.process_pdf(input_data, metadata)
                elif input_type == "text":
                    content, chunks = self.input_processor.process_text(input_data, metadata)
                elif input_type == "json":
                    content, chunks = self.input_processor.process_json(input_data, metadata)
                elif input_type == "docx":
                    content, chunks = self.input_processor.process_docx(input_data, metadata)
                elif input_type == "google_docs":
                    content, chunks = self.input_processor.process_google_docs(input_data, metadata)
                elif input_type == "wiki":
                    content, chunks = self.input_processor.process_wiki(input_data, metadata)
                elif input_type == "wiki_file":
                    content, chunks = self.input_processor.process_wiki_file(input_data, metadata)
                else:
                    raise ValueError(f"Unsupported input type: {input_type}. Supported types: pdf, text, json, docx, google_docs, wiki, wiki_file")
                
                # Check cache if enabled
                cache_key = f"{input_type}_{hashlib.md5(content.encode()).hexdigest()}"
                if self.cache:
                    cached_result = self.cache.get(content, self.config)
                    if cached_result:
                        logger.info("Using cached extraction results")
                        extraction_result = cached_result
                    else:
                        extraction_result = self._perform_extraction(content, chunks, questions)
                        self.cache.set(content, self.config, extraction_result)
                else:
                    extraction_result = self._perform_extraction(content, chunks, questions)
                
                # Create Document object
                # Ensure metadata is not empty for ChromaDB 0.6.3 compatibility
                document_metadata = metadata or {}
                if not document_metadata:
                    document_metadata = {"source": "document_upload", "created_by": created_by}
                
                document = self._create_document(
                    content=content,
                    source_type=source_type,
                    document_type=document_type,
                    created_by=created_by,
                    domain_id=domain_id,
                    metadata=document_metadata
                )
                
                # Create DocumentInsight object with chunks
                document_insight = self._create_document_insight(
                    document=document,
                    extraction_result=extraction_result,
                    source_type=source_type,
                    document_type=document_type,
                    event_type=event_type,
                    created_by=created_by,
                    domain_id=domain_id,
                    chunks=chunks
                )
                
                # Store in databases using session manager
                if self.config.store_in_db:
                    await self._store_in_database(session, document, document_insight)
                
                if self.config.store_in_chromadb and self.chroma_store:
                    chromadb_ids = self._store_in_chromadb(document, document_insight, chunks)
                    document_insight.chromadb_ids = chromadb_ids
                    if not chromadb_ids:
                        logger.warning(f"ChromaDB storage failed for document {document.id} - no IDs returned")
                    else:
                        logger.info(f"ChromaDB storage successful for document {document.id} - {len(chromadb_ids)} chunks stored")
                
                await session.commit()
                
                processing_time = (datetime.now() - start_time).total_seconds()
                logger.info(f"Document ingestion completed in {processing_time:.2f} seconds")
                
                return document, document_insight
                
            except Exception as e:
                await session.rollback()
                logger.error(f"Error in document ingestion: {e}")
                raise
    
    async def ingest_document_auto(self,
                                  input_data: Union[str, Dict],
                                  source_type: str,
                                  document_type: str,
                                  created_by: str,
                                  domain_id: str,
                                  questions: Optional[List[str]] = None,
                                  metadata: Optional[Dict] = None,
                                  event_type: str = "document_ingestion") -> Tuple[Document, DocumentInsight]:
        """
        Auto-detect document type and ingest document
        
        Args:
            input_data: File path or content
            source_type: Source of the document
            document_type: Type/category of document
            created_by: User/system that created the document
            domain_id: Domain ID for the document
            questions: Optional questions for targeted extraction
            metadata: Additional metadata
            event_type: Type of processing event
            
        Returns:
            Tuple of (Document, DocumentInsight)
        """
        detected_type = self._detect_document_type(input_data)
        return await self.ingest_document(
            input_data=input_data,
            input_type=detected_type,
            source_type=source_type,
            document_type=document_type,
            created_by=created_by,
            domain_id=domain_id,
            questions=questions,
            metadata=metadata,
            event_type=event_type
        )
    
    def _detect_document_type(self, input_data: Union[str, Dict]) -> str:
        """Auto-detect document type based on input data"""
        if isinstance(input_data, dict):
            # Check if it's a Google Docs reference
            if "google_doc_id" in input_data:
                return "google_docs"
            # Check if it's wiki content
            elif "wiki_content" in input_data:
                return "wiki"
            else:
                return "json"
        
        if isinstance(input_data, str):
            # Check file extension
            if input_data.endswith('.pdf'):
                return "pdf"
            elif input_data.endswith('.docx'):
                return "docx"
            elif input_data.endswith(('.md', '.wiki', '.txt')):
                return "wiki_file"
            # Check if it looks like Google Docs ID
            elif len(input_data) > 20 and input_data.isalnum():
                return "google_docs"
            # Check if it looks like wiki content
            elif any(marker in input_data for marker in ['[[', '{{', '==', "'''", "''"]):
                return "wiki"
            else:
                return "text"
        
        return "text"
    
    def _perform_extraction(self, content: str, chunks: List[DocumentChunk], questions: Optional[List[str]]) -> Dict:
        """Perform extraction based on configuration"""
        if questions and self.config.use_questions:
            return self.extractor.extract_with_questions(content, questions)
        else:
            return self.extractor.extract_basic_insights(content)
    
    def _create_document(self, content: str, source_type: str, document_type: str, 
                        created_by: str, domain_id: str, metadata: Dict) -> Document:
        """Create Document dataclass instance"""
        doc_id = uuid.uuid4()
        
        return Document(
            id=doc_id,
            document_id=doc_id,  # Using same ID for simplicity
            version=1,
            content=content,
            json_metadata=metadata,
            source_type=source_type,
            document_type=document_type,
            created_at=datetime.now(),
            created_by=created_by,
            domain_id=domain_id
        )
    
    def _create_document_insight(self, document: Document, extraction_result: Dict,
                                source_type: str, document_type: str, event_type: str,
                                created_by: str, domain_id: str, chunks: List[DocumentChunk] = None) -> DocumentInsight:
        """Create DocumentInsight dataclass instance with simplified structure"""
        
        current_time = datetime.now()
        extraction_date_iso = current_time.isoformat()
        
        # Extract insights and metadata from extraction result
        base_insights = extraction_result.get("insights", {})
        extraction_metadata = extraction_result.get("extraction_metadata", {})
        extraction_types = extraction_result.get("extraction_types", [])
        
        # Prepare chunk content and key phrases
        chunk_content = self._prepare_chunk_content(chunks) if chunks else ""
        key_phrases = self._extract_key_phrases(extraction_result, chunks)
        
        # Create additional insights structure for complex extractions
        additional_insights = {
            **base_insights,  # All extraction type results
            # Add legacy fields for backward compatibility
            "extracted_entities": extraction_result.get("entities", {}),
            "ner_text": extraction_result.get("ner_text", ""),
            "legacy_insights": extraction_result.get("insights", {}),
            # Add chunk metadata
            "chunks": self._prepare_chunk_data(chunks) if chunks else [],
            # Add document metadata
            "document_metadata": {
                "document_id": str(document.id),
                "source_type": source_type,
                "document_type": document_type,
                "created_by": created_by,
                "extraction_date": extraction_date_iso
            }
        }
        
        # Create extraction configuration for storage
        extraction_config = {
            "extraction_types": extraction_types,
            "extraction_metadata": extraction_metadata,
            "extraction_timestamp": extraction_date_iso,
            "extraction_version": "5.0",  # Version to track simplified extraction
            "model_name": self.config.model_name,
            "temperature": self.config.temperature,
            "chunk_count": len(chunks) if chunks else 0
        }
        
        return DocumentInsight(
            id=str(uuid.uuid4()),
            document_id=str(document.id),
            source_type=source_type,
            document_type=document_type,
            event_timestamp=extraction_date_iso,
            chromadb_ids=[],  # Will be populated later
            event_type=event_type,
            created_by=created_by,
            domain_id=domain_id,
            document=document,
            # Simplified structure
            chunk_content=chunk_content,
            key_phrases=key_phrases,
            # Additional insights for complex extractions
            insights=additional_insights,
            extraction_config=extraction_config,
            extraction_date=extraction_date_iso
        )
    
    def _prepare_chunk_content(self, chunks: List[DocumentChunk]) -> str:
        """Prepare consolidated chunk content as text"""
        if not chunks:
            return ""
        
        # Combine all chunk content into a single text
        content_parts = []
        for chunk in sorted(chunks, key=lambda x: x.metadata.get("chunk_index", 0)):
            content_parts.append(chunk.content)
        
        return "\n\n".join(content_parts)
    
    def _extract_key_phrases(self, extraction_result: Dict, chunks: List[DocumentChunk]) -> List[str]:
        """Extract key phrases from extraction results and chunk content"""
        key_phrases = set()
        
        # Extract phrases from extraction results
        if "phrases" in extraction_result:
            key_phrases.update(extraction_result["phrases"])
        
        # Extract phrases from business intelligence insights
        if "insights" in extraction_result and "business_intelligence" in extraction_result["insights"]:
            bi_insights = extraction_result["insights"]["business_intelligence"]
            if "phrases" in bi_insights:
                key_phrases.update(bi_insights["phrases"])
        
        # Extract key terms from other insight types
        if "insights" in extraction_result:
            for insight_type, data in extraction_result["insights"].items():
                if isinstance(data, dict):
                    # Extract key terms from various insight types
                    if "key_terms" in data:
                        key_phrases.update(data["key_terms"])
                    if "important_phrases" in data:
                        key_phrases.update(data["important_phrases"])
                    if "concepts" in data:
                        key_phrases.update(data["concepts"])
        
        # Extract entities as key phrases
        if "entities" in extraction_result:
            entities = extraction_result["entities"]
            if isinstance(entities, dict):
                for entity_type, entity_list in entities.items():
                    if isinstance(entity_list, list):
                        key_phrases.update(entity_list)
        
        # Extract from chunk content using simple keyword extraction
        if chunks:
            chunk_text = self._prepare_chunk_content(chunks)
            # Simple keyword extraction (can be enhanced with NLP libraries)
            words = chunk_text.lower().split()
            # Filter for potential key terms (longer words, capitalized words, etc.)
            potential_keywords = [
                word.strip(".,!?;:()[]{}") for word in words 
                if len(word) > 3 and (word[0].isupper() or word.isalpha())
            ]
            # Add most frequent terms
            from collections import Counter
            word_counts = Counter(potential_keywords)
            key_phrases.update([word for word, count in word_counts.most_common(10) if count > 1])
        
        # Clean and return unique phrases
        cleaned_phrases = []
        for phrase in key_phrases:
            if phrase and isinstance(phrase, str) and len(phrase.strip()) > 2:
                cleaned_phrases.append(phrase.strip())
        
        return list(set(cleaned_phrases))[:50]  # Limit to top 50 phrases
    
    def _prepare_chunk_data(self, chunks: List[DocumentChunk]) -> List[Dict]:
        """Prepare chunk data for storage in insights"""
        chunk_data = []
        for chunk in chunks:
            chunk_data.append({
                "chunk_id": chunk.chunk_id,
                "content": chunk.content,
                "metadata": chunk.metadata,
                "length": len(chunk.content),
                "index": chunk.metadata.get("chunk_index", 0)
            })
        return chunk_data
    
    async def _store_in_database(self, session: AsyncSession, document: Document, insight: DocumentInsight):
        """Store document and insight in PostgreSQL using session manager"""
        try:
            from app.schemas.docs.docmodels import DocumentVersion, DocumentInsightVersion
            
            # Create DocumentVersion record
            doc_version = DocumentVersion(
                id=document.id,
                document_id=document.document_id,
                version=document.version,
                content=document.content,
                json_metadata=document.json_metadata,
                source_type=document.source_type,
                document_type=document.document_type,
                created_at=document.created_at,
                created_by=document.created_by,
                domain_id=document.domain_id
            )
            session.add(doc_version)
            
            # Create DocumentInsightVersion record
            insight_version = DocumentInsightVersion(
                id=insight.id,
                insight_id=insight.id,
                document_id=document.id,  # Add missing document_id
                version=1,  # Default version for insights
                source_type=insight.source_type,
                document_type=insight.document_type,
                event_timestamp=datetime.fromisoformat(insight.event_timestamp.replace('Z', '+00:00')) if insight.event_timestamp else datetime.now(),
                chromadb_ids=insight.chromadb_ids,
                event_type=insight.event_type,
                created_at=datetime.now(),
                created_by=insight.created_by,
                domain_id=insight.domain_id,  # Add missing domain_id
                # Simplified structure
                chunk_content=insight.chunk_content,
                key_phrases=insight.key_phrases,
                # Additional insights for complex extractions
                insights=insight.insights,
                extraction_config=insight.extraction_config,
                extraction_date=datetime.fromisoformat(insight.extraction_date.replace('Z', '+00:00')) if insight.extraction_date else None
            )
            session.add(insight_version)
            
            logger.info(f"Stored document {document.id} with insights in database")
            logger.info(f"Chunk content length: {len(insight.chunk_content) if insight.chunk_content else 0}")
            logger.info(f"Key phrases count: {len(insight.key_phrases) if insight.key_phrases else 0}")
            
        except Exception as e:
            logger.error(f"Error storing in database: {e}")
            raise
    
    def _store_in_chromadb(self, document: Document, insight: DocumentInsight, 
                          chunks: List[DocumentChunk]) -> List[str]:
        """Store document chunks in ChromaDB with flexible insights metadata"""
        try:
            logger.info(f"Starting ChromaDB storage for document {document.id}")
            logger.info(f"Document type: {document.document_type}, source: {document.source_type}")
            logger.info(f"Number of chunks to process: {len(chunks) if chunks else 0}")
            logger.info(f"ChromaDB store initialized: {self.chroma_store is not None}")
            logger.info(f"TF-IDF enabled: {self.chroma_store.tf_idf if self.chroma_store else False}")
            
            chromadb_ids = []
            
            # Parse extraction date for metadata
            extraction_datetime = datetime.fromisoformat(insight.extraction_date.replace('Z', '+00:00')) if insight.extraction_date else datetime.now()
            
            # Create flexible metadata for filtering based on available insights
            base_metadata = {
                        "document_id": str(document.id),
                        "insight_id": insight.id,
                        "document_type": document.document_type,
                "source_type": document.source_type,
                "created_by": document.created_by,
                "extraction_date": insight.extraction_date,
                "event_type": insight.event_type,
                "event_timestamp": insight.event_timestamp,
                # Date metadata for filtering
                "extraction_year": extraction_datetime.year,
                "extraction_month": extraction_datetime.month,
                "extraction_day": extraction_datetime.day,
                # Extraction configuration metadata
                "extraction_types": insight.extraction_config.get("extraction_types", []) if insight.extraction_config else [],
                "extraction_version": insight.extraction_config.get("extraction_version", "1.0") if insight.extraction_config else "1.0",
            }
            
            logger.debug(f"Base metadata created with {len(base_metadata)} fields: {list(base_metadata.keys())}")
            
            # Add simplified metadata
            base_metadata.update({
                "chunk_content_length": len(insight.chunk_content) if insight.chunk_content else 0,
                "key_phrases_count": len(insight.key_phrases) if insight.key_phrases else 0,
                "has_chunk_content": bool(insight.chunk_content),
                "has_key_phrases": bool(insight.key_phrases),
                "key_phrases": insight.key_phrases or []
            })
            
            # Add dynamic metadata based on additional insights
            if insight.insights:
                # Add metadata for extraction types
                for extraction_type, data in insight.insights.items():
                    if isinstance(data, dict) and extraction_type not in ["chunks", "document_metadata", "extracted_entities", "ner_text", "phrases", "legacy_insights"]:
                        # Add has_* flags for each extraction type
                        base_metadata[f"has_{extraction_type}"] = bool(data and len(data) > 0)
                        # Add count metadata
                        base_metadata[f"{extraction_type}_count"] = len(data) if data else 0
                        
                        # Add specific metadata for business intelligence
                        if extraction_type == "business_intelligence":
                            base_metadata["has_kpis"] = bool(data.get("kpis"))
                            base_metadata["has_business_terms"] = bool(data.get("business_terms"))
                            base_metadata["has_data_definitions"] = bool(data.get("data_definitions"))
                            base_metadata["has_operational_metrics"] = bool(data.get("operational_metrics"))
                            base_metadata["kpi_count"] = len(data.get("kpis", {}))
                            base_metadata["business_term_count"] = len(data.get("business_terms", {}))
                            base_metadata["data_definition_count"] = len(data.get("data_definitions", {}))
                            base_metadata["operational_metric_count"] = len(data.get("operational_metrics", {}))
                
                # Add chunk metadata from insights
                chunks = insight.insights.get("chunks", [])
                base_metadata["chunk_count"] = len(chunks)
                base_metadata["has_chunks"] = bool(chunks)
                
                # Add legacy field metadata
                base_metadata["has_entities"] = bool(insight.insights.get("extracted_entities"))
                base_metadata["has_ner_text"] = bool(insight.insights.get("ner_text"))
                base_metadata["entity_count"] = len(insight.insights.get("extracted_entities", {}))
            
            # Create Langchain documents from chunks
            logger.info(f"Creating Langchain documents from {len(chunks)} chunks for document {document.id}")
            langchain_docs = []
            
            # Validate chunks before processing
            logger.debug(f"Chunks type: {type(chunks)}, length: {len(chunks) if chunks else 0}")
            
            for i, chunk in enumerate(chunks):
                try:
                    # Handle both DocumentChunk objects and dictionaries
                    if isinstance(chunk, dict):
                        # Extract data from dictionary format
                        chunk_content = chunk.get('content', chunk.get('page_content', ''))
                        chunk_meta = chunk.get('metadata', {})
                        chunk_id = chunk.get('chunk_id', f"chunk_{i}")
                    else:
                        # Handle DocumentChunk object format
                        if not hasattr(chunk, 'metadata'):
                            logger.error(f"Chunk {i} does not have metadata attribute. Chunk type: {type(chunk)}")
                            continue
                        
                        if not hasattr(chunk, 'content'):
                            logger.error(f"Chunk {i} does not have content attribute")
                            continue
                        
                        if not hasattr(chunk, 'chunk_id'):
                            logger.error(f"Chunk {i} does not have chunk_id attribute")
                            continue
                        
                        chunk_content = chunk.content
                        chunk_meta = chunk.metadata or {}
                        chunk_id = chunk.chunk_id
                    
                    chunk_metadata = {
                        **base_metadata,
                        **chunk_meta,
                        "chunk_index": i,
                        "chunk_id": chunk_id,
                        "chunk_length": len(chunk_content),
                        # Add all insights data to metadata for searchability
                        "insights": insight.insights or {},
                        "extraction_config": insight.extraction_config or {},
                    }
                    
                    # Ensure metadata is not empty for ChromaDB 0.6.3 compatibility
                    if not chunk_metadata:
                        logger.warning(f"Chunk {i} metadata was empty, using default metadata")
                        chunk_metadata = {"source": "document_upload", "chunk_index": i}
                    
                    doc = LangchainDocument(
                        page_content=chunk_content,
                        metadata=chunk_metadata
                    )
                    langchain_docs.append(doc)
                    
                except Exception as chunk_error:
                    logger.error(f"Error processing chunk {i}: {chunk_error}")
                    logger.error(f"Chunk {i} error type: {type(chunk_error).__name__}")
                    raise chunk_error
            
            logger.info(f"Created {len(langchain_docs)} Langchain documents, attempting ChromaDB storage")
            
            # Store in ChromaDB
            try:
                logger.info(f"Calling chroma_store.add_documents with {len(langchain_docs)} documents")
                ids = self.chroma_store.add_documents(langchain_docs)
                logger.info(f"ChromaDB storage completed. Returned: {ids}")
                logger.info(f"ChromaDB storage successful: {len(ids) if ids else 0} document IDs returned")
                if ids:
                    logger.debug(f"ChromaDB IDs: {ids[:3]}{'...' if len(ids) > 3 else ''}")
                else:
                    logger.warning("ChromaDB storage returned empty IDs list")
                    # Add more debugging to understand why no IDs are returned
                    logger.warning("ChromaDB storage completed but returned no IDs - this may indicate an issue with the ChromaDB wrapper")
            except Exception as chroma_error:
                logger.error(f"ChromaDB storage failed: {chroma_error}")
                logger.error(f"ChromaDB error type: {type(chroma_error).__name__}")
                logger.error(f"ChromaDB error details: {str(chroma_error)}")
                raise chroma_error
            
            # Store chunk content and key phrases for TF-IDF indexing
            if insight.chunk_content and self.chroma_store.tf_idf:
                logger.info(f"Starting TF-IDF storage for {len(langchain_docs)} documents")
                try:
                    self._store_tfidf_content(document, insight, langchain_docs)
                    logger.info("TF-IDF storage completed successfully")
                except Exception as tfidf_error:
                    logger.error(f"TF-IDF storage failed: {tfidf_error}")
                    # Don't raise here as TF-IDF is optional
            
            logger.info(f"Successfully stored {len(langchain_docs)} chunks in ChromaDB with flexible insights metadata")
            return ids or []
            
        except Exception as e:
            logger.error(f"Error storing in ChromaDB: {e}")
            return []
    
    def _store_tfidf_content(self, document: Document, insight: DocumentInsight, langchain_docs: List[LangchainDocument]):
        """Store chunk content and key phrases for TF-IDF indexing"""
        try:
            logger.info(f"Starting TF-IDF content storage for document {document.id}")
            logger.info(f"Chunk content length: {len(insight.chunk_content) if insight.chunk_content else 0}")
            logger.info(f"Key phrases count: {len(insight.key_phrases) if insight.key_phrases else 0}")
            logger.info(f"Langchain docs count: {len(langchain_docs)}")
            
            # Create documents for TF-IDF indexing
            tfidf_docs = []
            
            # Add consolidated chunk content as a single document
            if insight.chunk_content:
                chunk_content_doc = LangchainDocument(
                    page_content=insight.chunk_content,
                    metadata={
                        "document_id": str(document.id),
                        "insight_id": str(insight.id),
                        "content_type": "chunk_content",
                        "source_type": insight.source_type,
                        "document_type": insight.document_type,
                        "extraction_date": insight.extraction_date,
                        "content_length": len(insight.chunk_content),
                        "key_phrases_count": len(insight.key_phrases) if insight.key_phrases else 0
                    }
                )
                tfidf_docs.append(chunk_content_doc)
            
            # Add key phrases as a separate document for phrase-based search
            if insight.key_phrases:
                phrases_text = " ".join(insight.key_phrases)
                phrases_doc = LangchainDocument(
                    page_content=phrases_text,
                    metadata={
                        "document_id": str(document.id),
                        "insight_id": str(insight.id),
                        "content_type": "key_phrases",
                        "source_type": insight.source_type,
                        "document_type": insight.document_type,
                        "extraction_date": insight.extraction_date,
                        "phrases_count": len(insight.key_phrases),
                        "phrases": insight.key_phrases
                    }
                )
                tfidf_docs.append(phrases_doc)
            
            # Add individual chunk documents for detailed search
            for i, doc in enumerate(langchain_docs):
                chunk_tfidf_doc = LangchainDocument(
                    page_content=doc.page_content,
                    metadata={
                        **doc.metadata,
                        "content_type": "chunk",
                        "chunk_index": i
                    }
                )
                tfidf_docs.append(chunk_tfidf_doc)
            
            # Store TF-IDF documents
            if tfidf_docs:
                self.chroma_store.add_tfidf_vectors(tfidf_docs, [str(uuid.uuid4()) for _ in tfidf_docs])
                logger.info(f"Stored {len(tfidf_docs)} documents for TF-IDF indexing")
            
        except Exception as e:
            logger.error(f"Error storing TF-IDF content: {e}")
    
    def search_documents(self, query: str, document_type: Optional[str] = None,
                        k: int = 5, use_tfidf: bool = True) -> List[Dict]:
        """Search documents using semantic and/or TF-IDF search"""
        if not self.chroma_store:
            raise ValueError("ChromaDB store not initialized")
        
        where_filter = {}
        if document_type:
            where_filter["document_type"] = document_type
        
        if use_tfidf:
            return self.chroma_store.semantic_search_with_tfidf(
                query=query,
                k=k,
                where=where_filter
            )
        else:
            return self.chroma_store.semantic_search(
                query=query,
                k=k,
                where=where_filter
            )
    
    def search_by_key_phrases(self, key_phrases: List[str] = None,
                             query: str = None,
                             document_type: Optional[str] = None,
                             source_type: Optional[str] = None,
                             extraction_year: Optional[int] = None,
                             extraction_month: Optional[int] = None,
                             min_phrase_count: Optional[int] = None,
                             k: int = 5) -> List[Dict]:
        """Search documents by key phrases and other filters"""
        if not self.chroma_store:
            raise ValueError("ChromaDB store not initialized")
        
        # Build where filter with proper ChromaDB syntax
        where_conditions = []
        
        # Basic filters
        if document_type:
            where_conditions.append({"document_type": document_type})
        if source_type:
            where_conditions.append({"source_type": source_type})
        
        # Key phrase filters
        if key_phrases:
            where_conditions.append({"key_phrases": {"$in": key_phrases}})
        
        # Date filters
        if extraction_year is not None:
            where_conditions.append({"extraction_year": extraction_year})
        if extraction_month is not None:
            where_conditions.append({"extraction_month": extraction_month})
        
        # Count filters
        if min_phrase_count is not None:
            where_conditions.append({"key_phrases_count": {"$gte": min_phrase_count}})
        
        # Combine conditions with $and operator
        if where_conditions:
            if len(where_conditions) == 1:
                where_filter = where_conditions[0]
            else:
                where_filter = {"$and": where_conditions}
        else:
            where_filter = None
        
        if query:
            return self.chroma_store.semantic_search_with_tfidf(
                query=query,
                k=k,
                where=where_filter
            )
        else:
            # If no query, use a generic search with filters
            # ChromaDB doesn't have a direct metadata-only search, so we use a generic query
            return self.chroma_store.semantic_search(
                query="",  # Empty query to get all documents
                k=k,
                where=where_filter
            )
    
    def search_by_keywords_tfidf(self, keywords: List[str] = None,
                                query: str = None,
                                content_type: str = "chunk_content",  # "chunk_content", "key_phrases", "chunk"
                                document_type: Optional[str] = None,
                                source_type: Optional[str] = None,
                                k: int = 5) -> List[Dict]:
        """Search documents using TF-IDF for efficient keyword retrieval"""
        if not self.chroma_store or not self.chroma_store.tf_idf:
            raise ValueError("ChromaDB store with TF-IDF not initialized")
        
        # Build where filter with proper ChromaDB syntax
        where_conditions = []
        
        # Basic filters
        if document_type:
            where_conditions.append({"document_type": document_type})
        if source_type:
            where_conditions.append({"source_type": source_type})
        
        # Content type filter
        where_conditions.append({"content_type": content_type})
        
        # Combine conditions with $and operator
        if where_conditions:
            if len(where_conditions) == 1:
                where_filter = where_conditions[0]
            else:
                where_filter = {"$and": where_conditions}
        else:
            where_filter = None
        
        # Create search query
        if keywords and query:
            search_query = f"{query} {' '.join(keywords)}"
        elif keywords:
            search_query = " ".join(keywords)
        elif query:
            search_query = query
        else:
            raise ValueError("Either keywords or query must be provided")
        
        # Perform TF-IDF search
        return self.chroma_store.tfidf_search(
            query=search_query,
            k=k,
            where=where_filter
        )
    
    def search_chunk_content_tfidf(self, query: str, document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search consolidated chunk content using TF-IDF"""
        return self.search_by_keywords_tfidf(
            query=query,
            content_type="chunk_content",
            document_type=document_type,
            k=k
        )
    
    def search_key_phrases_tfidf(self, keywords: List[str], document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search key phrases using TF-IDF"""
        return self.search_by_keywords_tfidf(
            keywords=keywords,
            content_type="key_phrases",
            document_type=document_type,
            k=k
        )
    
    def search_chunks_tfidf(self, query: str, document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search individual chunks using TF-IDF"""
        return self.search_by_keywords_tfidf(
            query=query,
            content_type="chunk",
            document_type=document_type,
            k=k
        )
    
    def search_by_insights(self, query: str = None, 
                          extraction_types: List[str] = None,
                          has_insight_type: str = None,
                          document_type: Optional[str] = None,
                          source_type: Optional[str] = None,
                          extraction_year: Optional[int] = None,
                          extraction_month: Optional[int] = None,
                          min_insight_count: Optional[int] = None,
                          k: int = 5) -> List[Dict]:
        """Search documents by flexible insights metadata filters"""
        if not self.chroma_store:
            raise ValueError("ChromaDB store not initialized")
        
        # Build where filter with proper ChromaDB syntax
        where_conditions = []
        
        # Basic filters
        if document_type:
            where_conditions.append({"document_type": document_type})
        if source_type:
            where_conditions.append({"source_type": source_type})
        
        # Extraction type filters
        if extraction_types:
            where_conditions.append({"extraction_types": {"$in": extraction_types}})
        
        # Specific insight type filters
        if has_insight_type:
            where_conditions.append({f"has_{has_insight_type}": True})
        
        # Date filters
        if extraction_year is not None:
            where_conditions.append({"extraction_year": extraction_year})
        if extraction_month is not None:
            where_conditions.append({"extraction_month": extraction_month})
        
        # Count filters
        if min_insight_count is not None and has_insight_type:
            where_conditions.append({f"{has_insight_type}_count": {"$gte": min_insight_count}})
        
        # Combine conditions with $and operator
        if where_conditions:
            if len(where_conditions) == 1:
                where_filter = where_conditions[0]
            else:
                where_filter = {"$and": where_conditions}
        else:
            where_filter = None
        
        if query:
            return self.chroma_store.semantic_search_with_tfidf(
                query=query,
                k=k,
                where=where_filter
            )
        else:
            # If no query, use a generic search with filters
            # ChromaDB doesn't have a direct metadata-only search, so we use a generic query
            return self.chroma_store.semantic_search(
                query="",  # Empty query to get all documents
                k=k,
                where=where_filter
            )
    
    def search_insight_type(self, insight_type: str, query: str = None, 
                           document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search for documents containing specific insight type"""
        return self.search_by_insights(
            query=query,
            has_insight_type=insight_type,
            document_type=document_type,
            k=k
        )
    
    def search_business_intelligence(self, query: str = None, document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search for documents containing business intelligence insights"""
        return self.search_insight_type("business_intelligence", query, document_type, k)
    
    def search_entities(self, query: str = None, document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search for documents containing entity insights"""
        return self.search_insight_type("entities", query, document_type, k)
    
    def search_financial_metrics(self, query: str = None, document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search for documents containing financial metrics"""
        return self.search_insight_type("financial_metrics", query, document_type, k)
    
    def search_compliance_terms(self, query: str = None, document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search for documents containing compliance terms"""
        return self.search_insight_type("compliance_terms", query, document_type, k)
    
    # Legacy methods for backward compatibility
    def search_kpis(self, query: str = None, document_type: Optional[str] = None, k: int = 5) -> List[Dict]:
        """Search for documents containing KPIs (legacy method)"""
        return self.search_by_insights(
            query=query,
            has_insight_type="business_intelligence",
            document_type=document_type,
            k=k
            )
    
    def get_document_by_id(self, document_id: str) -> Optional[Document]:
        """Retrieve document by ID from database"""
        # Note: This method should be used through DocumentPersistenceService
        # for proper session management
        logger.warning("get_document_by_id should be used through DocumentPersistenceService")
        return None
    
    def get_insights_by_document_id(self, document_id: str) -> List[DocumentInsight]:
        """Retrieve all insights for a document"""
        # Note: This method should be used through DocumentPersistenceService
        # for proper session management
        logger.warning("get_insights_by_document_id should be used through DocumentPersistenceService")
        return []
    
    def get_insights_by_type(self, document_id: str, insight_type: str) -> Dict:
        """Retrieve specific insight type from a document"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0].insights:
            return insights[0].insights.get(insight_type, {})
        return {}
    
    def get_all_insights(self, document_id: str) -> Dict:
        """Retrieve all insights from a document"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0]:
            insight = insights[0]
            return {
                "insights": insight.insights or {},
                "extraction_config": insight.extraction_config or {},
                "extraction_date": insight.extraction_date,
                "extraction_types": insight.extraction_config.get("extraction_types", []) if insight.extraction_config else []
            }
        return {}
    
    def get_extraction_config(self, document_id: str) -> Dict:
        """Retrieve extraction configuration from a document"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0].extraction_config:
            return insights[0].extraction_config
        return {}
    
    def get_chunk_content(self, document_id: str) -> str:
        """Retrieve chunk content from a document"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0].chunk_content:
            return insights[0].chunk_content
        return ""
    
    def get_key_phrases(self, document_id: str) -> List[str]:
        """Retrieve key phrases from a document"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0].key_phrases:
            return insights[0].key_phrases
        return []
    
    def get_chunks(self, document_id: str) -> List[Dict]:
        """Retrieve chunk data from a document (from insights)"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0].insights:
            return insights[0].insights.get("chunks", [])
        return []
    
    def get_entities(self, document_id: str) -> Dict:
        """Retrieve extracted entities from a document"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0].insights:
            return insights[0].insights.get("extracted_entities", {})
        return {}
    
    def get_ner_text(self, document_id: str) -> str:
        """Retrieve NER text from a document"""
        insights = self.get_insights_by_document_id(document_id)
        if insights and insights[0].insights:
            return insights[0].insights.get("ner_text", "")
        return ""
    
    # Legacy methods for backward compatibility
    def get_kpis_by_document_id(self, document_id: str) -> Dict:
        """Retrieve KPIs extracted from a document (legacy method)"""
        business_intelligence = self.get_insights_by_type(document_id, "business_intelligence")
        return business_intelligence.get("kpis", {})
    
    def get_business_terms_by_document_id(self, document_id: str) -> Dict:
        """Retrieve business terms and definitions from a document (legacy method)"""
        business_intelligence = self.get_insights_by_type(document_id, "business_intelligence")
        return business_intelligence.get("business_terms", {})
    
    def get_data_definitions_by_document_id(self, document_id: str) -> Dict:
        """Retrieve data definitions and schema from a document (legacy method)"""
        business_intelligence = self.get_insights_by_type(document_id, "business_intelligence")
        return business_intelligence.get("data_definitions", {})
    
    def get_operational_metrics_by_document_id(self, document_id: str) -> Dict:
        """Retrieve operational metrics from a document (legacy method)"""
        business_intelligence = self.get_insights_by_type(document_id, "business_intelligence")
        return business_intelligence.get("operational_metrics", {})
    
    def get_all_business_intelligence(self, document_id: str) -> Dict:
        """Retrieve all business intelligence data from a document (legacy method)"""
        return self.get_insights_by_type(document_id, "business_intelligence")
    
    def get_extracted_entities(self, document_id: str) -> Dict:
        """Retrieve extracted entities from a document (legacy method)"""
        return self.get_entities(document_id)


# Helper function to create services with proper initialization
def create_services_with_config(config: ServiceConfig = None, chroma_client: chromadb.PersistentClient = None, chroma_path: str = "./chroma_db") -> Tuple['DocumentIngestionService', 'DocumentPersistenceService']:
    """Create both DocumentIngestionService and DocumentPersistenceService with proper initialization"""
    from app.service.document_persistence_service import create_document_persistence_service
    
    if config is None:
        config = ServiceConfig()
    
    # Initialize session manager
    session_manager = SessionManager(config)
    
    # Initialize domain manager (with None for async usage)
    domain_manager = DomainManager(None)
    
    # Create ingestion service with provided client or fallback to path
    ingestion_service = create_ingestion_service(
        session_manager=session_manager,
        domain_manager=domain_manager,
        chroma_client=chroma_client,
        chroma_path=chroma_path
    )
    
    # Create persistence service
    persistence_service = create_document_persistence_service(
        session_manager=session_manager,
        domain_manager=domain_manager
    )
    
    return ingestion_service, persistence_service

# Example usage and factory functions
def create_ingestion_service(
    session_manager: SessionManager,
    domain_manager: DomainManager,
    chroma_client: Optional[chromadb.PersistentClient] = None,
    chroma_path: str = "./chroma_db",
    collection_name: str = "documents",
    extraction_types: List[str] = None,
    custom_extraction_config: Dict = None,
    document_store: Any = None,
) -> DocumentIngestionService:
    """Factory function to create a configured ingestion service with session manager"""

    if document_store is None and chroma_client is None:
        from app.core.settings import get_settings, VectorStoreType

        if get_settings().VECTOR_STORE_TYPE == VectorStoreType.CHROMA:
            chroma_client = chromadb.PersistentClient(path=chroma_path)

    # Configure processing with flexible extraction types
    config = ProcessingConfig(
        chunk_size=2000,
        chunk_overlap=200,
        use_questions=True,
        enable_caching=True,
        store_in_db=True,
        store_in_chromadb=True,
        extraction_types=extraction_types or ["business_intelligence", "entities"],
        custom_extraction_config=custom_extraction_config
    )
    
    return DocumentIngestionService(
        session_manager=session_manager,
        domain_manager=domain_manager,
        config=config,
        chroma_client=chroma_client,
        document_store=document_store,
    )


def create_custom_ingestion_service(session_manager: SessionManager,
                                  domain_manager: DomainManager,
                                  extraction_configs: Dict[str, ExtractionConfig],
                                  chroma_path: str = "./chroma_db") -> DocumentIngestionService:
    """Create ingestion service with custom extraction configurations using session manager"""
    
    # Initialize ChromaDB client
    chroma_client = chromadb.PersistentClient(path=chroma_path)
    
    # Configure processing with custom extraction configs
    config = ProcessingConfig(
        chunk_size=2000,
        chunk_overlap=200,
        use_questions=True,
        enable_caching=True,
        store_in_db=True,
        store_in_chromadb=True,
        extraction_types=list(extraction_configs.keys()),
        custom_extraction_config=extraction_configs
    )
    
    return DocumentIngestionService(
        session_manager=session_manager,
        domain_manager=domain_manager,
        config=config,
        chroma_client=chroma_client
    )


# API Interface Example
class DocumentAPI:
    """REST API interface for the document ingestion service"""
    
    def __init__(self, service: DocumentIngestionService):
        self.service = service
    
    def ingest_pdf(self, file_path: str, metadata: Dict = None, questions: List[str] = None):
        """API endpoint for PDF ingestion"""
        return self.service.ingest_document(
            input_data=file_path,
            input_type="pdf",
            source_type="upload",
            document_type="pdf_document",
            created_by="api_user",
            questions=questions,
            metadata=metadata
        )
    
    def ingest_text(self, content: str, metadata: Dict = None, questions: List[str] = None):
        """API endpoint for text ingestion"""
        return self.service.ingest_document(
            input_data=content,
            input_type="text",
            source_type="direct_input",
            document_type="text_document",
            created_by="api_user",
            questions=questions,
            metadata=metadata
        )
    
    def ingest_json(self, json_data: Dict, metadata: Dict = None, questions: List[str] = None):
        """API endpoint for JSON ingestion"""
        return self.service.ingest_document(
            input_data=json_data,
            input_type="json",
            source_type="structured_data",
            document_type="json_document",
            created_by="api_user",
            questions=questions,
            metadata=metadata
        )
    
    def search(self, query: str, document_type: str = None, limit: int = 5):
        """API endpoint for document search"""
        return self.service.search_documents(
            query=query,
            document_type=document_type,
            k=limit,
            use_tfidf=True
        )


if __name__ == "__main__":
    import asyncio
    from app.core.settings import ServiceConfig
    
    async def main():
        # Example 1: Basic usage with different document types using session manager
        # Initialize services with proper configuration
        config = ServiceConfig()
        service, persistence_service = create_services_with_config(
            config=config,
            chroma_path="/Users/sameerm/ComplianceSpark/byziplatform/unstructured/docinsights/chroma_db"
        )
        
        # PDF document
        document, insight = await service.ingest_document(
        input_data="/Users/sameerm/Downloads/LEXYAI/p1_v1_Rexel_Case_Study_EN.pdf",
        input_type="pdf",
        source_type="upload",
        document_type="Revel LMS report",
        created_by="user123",
            domain_id="domain_123",
            questions=["What are the key financial metrics?", "What operational KPIs are tracked?"]
        )
    
        """ 
        # Word document
        word_doc, word_insight = await service.ingest_document(
            input_data="business_plan.docx",
            input_type="docx",
            source_type="file_upload",
            document_type="business_document",
            created_by="user123",
            domain_id="domain_123"
        )
        
        # Google Docs document
        google_doc, google_insight = await service.ingest_document(
            input_data="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgvE2upms",  # Google Docs ID
            input_type="google_docs",
            source_type="google_docs",
            document_type="collaborative_document",
            created_by="user123",
            domain_id="domain_123"
        )
        
        # Wiki content
        wiki_doc, wiki_insight = await service.ingest_document(
            input_data="company_wiki.md",
            input_type="wiki_file",
            source_type="wiki",
            document_type="knowledge_base",
            created_by="user123",
            domain_id="domain_123"
        )
        
        # Auto-detect document type
        auto_doc, auto_insight = await service.ingest_document_auto(
            input_data="unknown_document.docx",  # Will auto-detect as docx
            source_type="file_upload",
            document_type="document",
            created_by="user123",
            domain_id="domain_123"
        )
        """
        print(f"Document created: {document.id}")
        print(f"Insight created: {insight.id}")
        print(f"Extraction types: {insight.extraction_config.get('extraction_types', [])}")
        
        # Example 2: Access simplified structure using DocumentPersistenceService
        # Get main content and key phrases
        insights = await persistence_service.get_insights_by_document_id(str(document.id))
        if insights:
            latest_insight = insights[0]
            chunk_content = latest_insight.chunk_content or ""
            key_phrases = latest_insight.key_phrases or []
            
            print(f"Chunk content length: {len(chunk_content)}")
            print(f"Key phrases: {key_phrases[:10]}...")  # Show first 10 phrases
            
            # Access additional insights for complex extractions
            if latest_insight.insights:
                print(f"Available insight types: {list(latest_insight.insights.keys())}")
                
                # Access specific insight types
                business_intelligence = latest_insight.insights.get("business_intelligence", {})
                entities = latest_insight.insights.get("entities", {})
                
                print(f"Business intelligence insights: {len(business_intelligence)}")
                print(f"Entity insights: {len(entities)}")
                
                # Access detailed chunk data from insights
                chunks = latest_insight.insights.get("chunks", [])
                print(f"Document chunks: {len(chunks)}")
                if chunks:
                    print(f"First chunk content preview: {chunks[0].get('content', '')[:100]}...")
                
                # Access legacy fields from insights
                extracted_entities = latest_insight.insights.get("extracted_entities", {})
                ner_text = latest_insight.insights.get("ner_text", "")
                
                print(f"Extracted entities: {len(extracted_entities)}")
                print(f"NER text length: {len(ner_text)}")
        else:
            print("No insights found for document")
        
        # Example 3: Custom extraction configuration
        custom_configs = {
            "financial_metrics": ExtractionConfig(
                extraction_type="financial_metrics",
                enabled=True,
                priority=1,
                custom_prompt="Extract detailed financial metrics and ratios",
                output_format="json"
            ),
            "compliance_terms": ExtractionConfig(
                extraction_type="compliance_terms",
                enabled=True,
                priority=2,
                custom_prompt="Extract compliance and regulatory terms",
                output_format="json"
            )
        }
        
        # Create custom service with the same session manager and domain manager
        custom_service = create_custom_ingestion_service(
            session_manager=service.session_manager,
            domain_manager=service.domain_manager,
            extraction_configs=custom_configs
        )
        
        # Example 4: TF-IDF keyword search
        # Search consolidated chunk content using TF-IDF
        content_results = service.search_chunk_content_tfidf(
            query="revenue growth quarterly performance",
            document_type="financial_report"
        )
        print(f"Found {len(content_results)} documents with TF-IDF content search")
        
        # Search key phrases using TF-IDF
        phrase_results = service.search_key_phrases_tfidf(
            keywords=["revenue", "growth", "KPI", "financial"],
            document_type="financial_report"
        )
        print(f"Found {len(phrase_results)} documents with TF-IDF phrase search")
        
        # Search individual chunks using TF-IDF
        chunk_results = service.search_chunks_tfidf(
            query="quarterly performance metrics",
            document_type="financial_report"
        )
        print(f"Found {len(chunk_results)} chunks with TF-IDF search")
        
        # Example 5: Search by key phrases (metadata-based)
        # Search for documents containing specific key phrases
        phrase_metadata_results = service.search_by_key_phrases(
            key_phrases=["revenue", "growth", "KPI"],
            document_type="financial_report"
        )
        print(f"Found {len(phrase_metadata_results)} documents with key phrase metadata")
        
        # Search with query and key phrase filters
        query_results = service.search_by_key_phrases(
            query="quarterly performance",
            key_phrases=["financial", "metrics"],
            min_phrase_count=3
        )
        print(f"Found {len(query_results)} documents with query and phrase filters")
        
        # Example 6: Search with flexible insights
        # Search for documents with specific insight types
        bi_docs = service.search_business_intelligence("revenue growth", document_type="financial_report")
        entity_docs = service.search_entities("company names", document_type="financial_report")
        
        print(f"Found {len(bi_docs)} documents with business intelligence")
        print(f"Found {len(entity_docs)} documents with entities")
        
        # Example 7: Advanced filtering by insight types
        results = service.search_by_insights(
            query="quarterly performance",
            extraction_types=["business_intelligence", "financial_metrics"],
            document_type="financial_report",
            extraction_year=2024
        )
        print(f"Found {len(results)} documents with specified insight types")
        
        # Example 8: Legacy compatibility
        kpis = service.get_kpis_by_document_id(str(document.id))
        business_terms = service.get_business_terms_by_document_id(str(document.id))
        
        print(f"Legacy KPI access: {len(kpis)} KPIs found")
        print(f"Legacy business terms access: {len(business_terms)} terms found")
    
    # Run the async main function
    asyncio.run(main())